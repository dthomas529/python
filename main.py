# Program that creates a resume
from docx import Document
from docx.shared import Inches

document = Document()

# profile picture
# document.add_picture('image.png', width=Inches(2.5))

# name, phone number, and email details
name = input('What is your name? ')
phone_number = input('What is your cell phone number? ')
email = input('What is your email address? ')

document.add_paragraph(name + ' | ' + phone_number + ' | ' + email)

# about me
document.add_heading('About Me')
about_me = input('Tell me about yourself ')
document.add_paragraph(about_me)

# work experience
document.add_heading('Work Experience')
p = document.add_paragraph()

company = input('Enter company name ')
from_date = input('From Date ')
to_date = input('To Date ')

p.add_run(company + ' ').bold = True
p.add_run(from_date + '-' + to_date + '\n').italic = True

experience_details = input('Describe your experience at ' + company + ' ')
p.add_run(experience_details)

# more experiences
while True:
  has_more_experiences = input('Do you have more experiences? Yes or No ')
  if has_more_experiences.lower() == 'yes':
    p = document.add_paragraph()
    company = input('Enter company name ')
    from_date = input('From Date ')
    to_date = input('To Date')

    p.add_run(company + ' ').bold = True
    p.add_run(from_date + '-' + to_date + '\n').italic = True

    experience_details = input('Describe your experience at ' + company + ' ')
    p.add_run(experience_details)
  else:
    break

# add list of skills with bullet marks
  document.add_heading('Skills')
  skill = input('Enter your skill ')
  p = document.add_paragraph()
  p.style = 'List Bullet'

  has_another_skill = input('Do you have another skill? Yes or No ')
  if has_another_skill.lower() == 'yes':
    skill = input('Enter your skill ')
    p = document.add_paragraph()
    p.style = 'List Bullet'
  else:
    break

# footer
  section = document.section[0]
  footer = section.footer
  p = footer.paragraphs[0]
  p.text = 'CV generated using Python'

document.save('cv.docx')
